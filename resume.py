"""Resume parsing, tailoring, and DOCX export."""

import os, re, io, tempfile, copy as _copy
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st

from utils import DEEPSEEK_KEY, call_deepseek


def parse_resume(uploaded_file):
    ext = os.path.splitext(uploaded_file.name)[1].lower()
    if ext == '.pdf':
        from PyPDF2 import PdfReader
        reader = PdfReader(uploaded_file)
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if ext == '.docx':
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
    if ext == '.doc':
        return _try_read_doc(uploaded_file)
    return uploaded_file.read().decode("utf-8", errors="ignore").strip()


def _try_read_doc(uploaded_file):
    content = uploaded_file.read()
    common_words = ["the", "and", "for", "with", "experience", "years", "work", "team", "company"]
    text = content.decode("utf-8", errors="ignore")
    word_count = sum(1 for w in common_words if w in text.lower())
    if word_count >= 3:
        lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 3 and not all(c < ' ' or c > '~' or c == '\x00' for c in l.strip())]
        result = '\n'.join(lines)
        if len(result) > 50:
            return result
    tmp_path = None
    try:
        import subprocess
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tf:
            tmp_path = tf.name
            tf.write(content)
            tf.flush()
        out = subprocess.check_output(['antiword', tmp_path], timeout=10)
        return out.decode("utf-8", errors="ignore").strip()
    except Exception:
        pass
    finally:
        if tmp_path:
            try:
                os.unlink(tmp_path)
            except OSError:
                pass
    lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 3]
    result = '\n'.join(lines)
    return result if len(result) > 50 else ""


def tailor_resume_simple(resume, job):
    job_text = f"{job['title']} {job['description']} {job['requirements']}".lower()
    stop_words = {'this','that','with','from','your','have','will','they','about','their','would','which','the','and','for','are','but','not','you','all','can','had','has','was','were','been','being','more','some','than','then','its','also','what','when','where','who','how','our'}
    job_words = set(re.findall(r'\b[a-z]{4,}\b', job_text)) - stop_words
    resume_words = set(re.findall(r'\b[a-z]{4,}\b', resume.lower()))
    missing = sorted(job_words - resume_words)
    rate = len(job_words & resume_words) / max(len(job_words), 1) * 100
    return rate, missing[:20]


def tailor_resume_ai(resume, job):
    prompt = f"""Job: {job['title']} at {job['company']}
Description: {job.get('description','')[:2000]}
Requirements: {job.get('requirements','')[:1000]}

My current resume:
{resume[:3000]}

Analyze my resume against this job. Return:
1. Match percentage (estimate)
2. Top 5 missing skills/keywords I should add
3. 3 bullet points to add to my resume
4. Brief summary of my strengths
Keep it concise."""
    return call_deepseek(prompt, "You are an expert resume reviewer. Be honest and specific.")


def generate_tailored_resume(resume, job):
    sections = ["NAME: [Full Name from resume]"]
    if st.session_state.get("sec_summary", True):
        sections.append("SUMMARY: [2-3 sentence professional summary tailored to this role]")
    if st.session_state.get("sec_skills", True):
        sections.append("SKILLS: [comma-separated skills, prioritize those matching the job]")
    if st.session_state.get("sec_experience", True):
        sections.append("EXPERIENCE: [keep same roles, rewrite bullets to match job keywords]")
    if st.session_state.get("sec_education", True):
        sections.append("EDUCATION: [keep as-is]")
    if st.session_state.get("sec_certs", False):
        sections.append("CERTIFICATIONS: [keep as-is, if any]")

    prompt = f"""Job Title: {job['title']}
Company: {job['company']}
Job Description: {job.get('description','')[:2000]}
Requirements: {job.get('requirements','')[:1000]}

My current resume:
{resume[:3000]}

Rewrite my resume for this job. Rules:
- NEVER add anything I didn't actually do — no invented skills, metrics, or achievements
- Keep my real job titles, dates, and company names exactly as-is
- Use UK English spelling throughout: -ise not -ize, -our not -or, -re not -er, "travelling" not "traveling"
- Use strong action verbs (e.g. "Led", "Designed", "Built", "Optimised")
- Write like a human: "I built the API" not "Spearheaded the development of a RESTful API"
- Avoid semicolons and run-on sentences — keep each bullet to one clear idea
- Keep bullets concise — one line each when possible
- Weave ATS keywords from the job description naturally, using the employer's exact terminology (e.g. "CI/CD pipelines" not "deployment automation"), but apply UK spelling to all other words
Include ONLY these sections in this exact order:

{chr(10).join(sections)}

IMPORTANT: Place the most relevant keywords early in SUMMARY and SKILLS. Return the full rewritten resume with those exact headers."""
    return call_deepseek(prompt, "You are an expert resume writer. Use UK English spelling throughout (-ise not -ize, -our not -or, -re not -er). Use strong professional action verbs. Avoid semicolons and run-on sentences. Weave ATS keywords naturally. NEVER add anything not in the original resume.", max_tokens=1500)


def _replace_para_text(para, new_text):
    if not para.runs:
        para.add_run(new_text)
        return
    for run in para.runs:
        run.text = ""
    para.runs[0].text = new_text


def _clear_para(para):
    for run in para.runs:
        run.text = ""


def build_docx(tailored_text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    section_headers = {"SUMMARY:", "SKILLS:", "EXPERIENCE:", "EDUCATION:", "CERTIFICATIONS:"}
    lines = tailored_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("NAME:"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("NAME:", "").strip())
            run.bold = True
            run.font.size = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif any(line.startswith(h) for h in section_headers):
            p = doc.add_paragraph()
            parts = line.split(":", 1)
            run = p.add_run(parts[0] + ":")
            run.bold = True
            run.font.size = Pt(13)
            if len(parts) > 1 and parts[1].strip():
                p.add_run("  " + parts[1].strip())
        elif line.startswith("- ") or line.startswith("• "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(2)
            text = line.lstrip("-• ").strip()
            p.add_run("• " + text)
        else:
            p = doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_tailored_docx(original_bytes, tailored_text):
    """Edit the original DOCX in-place with tailored content, preserving all formatting."""
    try:
        from docx import Document as DocxDocument
        from io import BytesIO

        doc = DocxDocument(BytesIO(original_bytes))

        sections = {}
        current_section = "OTHER"
        for line in tailored_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith("NAME:"):
                sections["NAME"] = line.replace("NAME:", "").strip()
            elif any(line.startswith(h) for h in ["SUMMARY:", "SKILLS:", "EXPERIENCE:", "EDUCATION:", "CERTIFICATIONS:"]):
                parts = line.split(":", 1)
                current_section = parts[0]
                sections[current_section] = parts[1].strip() if len(parts) > 1 else ""
            else:
                sections[current_section] = sections.get(current_section, "") + "\n" + line

        SECTION_ORDER = ["NAME", "SUMMARY", "SKILLS", "EXPERIENCE", "EDUCATION", "CERTIFICATIONS"]
        sec_groups = {}
        current_sec = None
        pending_body = []

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            upper = text.upper()
            matched = None
            for sec in SECTION_ORDER:
                if upper.startswith(sec + ":") or upper == sec:
                    matched = sec
                    break
            if matched:
                if current_sec and current_sec in sec_groups:
                    sec_groups[current_sec]["body_idxs"] = pending_body
                current_sec = matched
                pending_body = []
                sec_groups[matched] = {"header_idx": i, "body_idxs": []}
            elif current_sec:
                pending_body.append(i)
        if current_sec and current_sec in sec_groups:
            sec_groups[current_sec]["body_idxs"] = pending_body

        for sec in reversed(SECTION_ORDER):
            if sec not in sections or sec not in sec_groups:
                continue
            group = sec_groups[sec]
            content = sections[sec]
            content_lines = [l.strip() for l in content.split('\n') if l.strip()]
            header_para = doc.paragraphs[group["header_idx"]]
            if sec == "NAME":
                _replace_para_text(header_para, f"NAME: {content}" if content else "NAME:")
            else:
                _replace_para_text(header_para, f"{sec}:")
            body_idxs = group["body_idxs"]
            for j, line in enumerate(content_lines):
                if j < len(body_idxs):
                    _replace_para_text(doc.paragraphs[body_idxs[j]], line)
                else:
                    template = doc.paragraphs[body_idxs[-1]] if body_idxs else header_para
                    clone_elem = _copy.deepcopy(template._element)
                    for r_elem in clone_elem.iter():
                        if r_elem.tag.endswith('}t'):
                            r_elem.text = ""
                    first_t = clone_elem.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')
                    if first_t is not None:
                        first_t.text = line
                    template._element.addnext(clone_elem)
            for j in range(len(content_lines), len(body_idxs)):
                _clear_para(doc.paragraphs[body_idxs[j]])

        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except Exception:
        return None
