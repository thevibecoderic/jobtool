"""Generate a sample resume .docx for testing."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# Name
p = doc.add_paragraph()
run = p.add_run("Alex Tan")
run.bold = True
run.font.size = Pt(18)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.add_run("alex.tan@email.com | +65 9123 4567 | linkedin.com/in/alextan").font.size = Pt(10)

doc.add_paragraph()

# Summary
doc.add_paragraph().add_run("SUMMARY").bold = True
doc.add_paragraph("Software engineer with 4 years of experience building web applications and APIs. Strong background in Python, JavaScript, and cloud infrastructure. Passionate about clean code and scalable architecture.")

# Skills
doc.add_paragraph().add_run("SKILLS").bold = True
doc.add_paragraph("Languages: Python, JavaScript, SQL, HTML/CSS")
doc.add_paragraph("Frameworks: React, Flask, FastAPI, Node.js")
doc.add_paragraph("Tools: Docker, Git, AWS (EC2, S3, Lambda), PostgreSQL, Redis")

# Experience
doc.add_paragraph().add_run("EXPERIENCE").bold = True
p = doc.add_paragraph()
p.add_run("Software Engineer  |  TechCorp Singapore  |  2022 – Present").bold = True
for b in [
    "Built REST APIs serving 10k+ daily users using FastAPI and PostgreSQL",
    "Migrated legacy monolith to microservices, reducing deployment time by 60%",
    "Implemented CI/CD pipelines with GitHub Actions and Docker",
    "Mentored 2 junior developers on Python best practices",
]:
    doc.add_paragraph(b, style='List Bullet')

p = doc.add_paragraph()
p.add_run("Junior Developer  |  StartUp Labs  |  2020 – 2022").bold = True
for b in [
    "Developed internal dashboard with React and Flask backend",
    "Wrote unit tests achieving 85% code coverage",
    "Managed AWS infrastructure including EC2 and S3",
]:
    doc.add_paragraph(b, style='List Bullet')

# Education
doc.add_paragraph().add_run("EDUCATION").bold = True
doc.add_paragraph("B.Sc. Computer Science  |  National University of Singapore  |  2016 – 2020")

doc.save("jobtool/sample_resume.docx")
print("✓ Created jobtool/sample_resume.docx")
