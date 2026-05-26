#!/usr/bin/env python3
"""
Job Scraper UI — Streamlit
Usage: streamlit run ui.py
Reads DEEPSEEK_API_KEY from .env or st.secrets
"""

import streamlit as st
import requests, re, json, os, time, urllib.parse, io, tempfile
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

st.set_page_config(page_title="Job Scraper", page_icon="🎯", layout="wide")

# Hide Streamlit branding and GitHub links (keeping Fork + Record buttons visible)
st.markdown("""
<style>
header, footer,
[data-testid="stFooter"],
.viewerBadge_container__1QSob,
a[href*="/creators/"], a[href*="github"],
iframe[src*="github"], div:has(> a[href*="github"])
{display:none !important;}
</style>
""", unsafe_allow_html=True)

# Force-remove viewer badge via DOM (catches elements CSS misses)
st.markdown("""
<div id="_x_hide_branding" style="display:none;"></div>
""", unsafe_allow_html=True)

# ── Config ────────────────────────────────────────────
HEADERS = {
    "User-Agent": "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/125.0.0.0 Safari/537.36",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}
GOOGLEBOT_HEADERS = {
    "User-Agent": "Mozilla/5.0 (compatible; Googlebot/2.1; +http://www.google.com/bot.html)",
    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    "Accept-Language": "en-US,en;q=0.9",
}
LOCATION = "Singapore"
TIME_RANGE = "r1209600"

def _load_env():
    if os.environ.get("DEEPSEEK_API_KEY"):
        return
    # Streamlit Community Cloud secrets
    try:
        import streamlit as _st
        if _st.secrets.get("DEEPSEEK_API_KEY"):
            os.environ["DEEPSEEK_API_KEY"] = _st.secrets["DEEPSEEK_API_KEY"]
            return
    except Exception:
        pass
    # Local .env fallback
    script_dir = os.path.dirname(os.path.abspath(__file__))
    env_path = os.path.join(script_dir, ".env")
    try:
        with open(env_path) as f:
            for line in f:
                line = line.strip()
                if line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                k, v = k.strip(), v.strip().strip('"').strip("'")
                if k == "DEEPSEEK_API_KEY":
                    os.environ["DEEPSEEK_API_KEY"] = v
    except FileNotFoundError:
        pass

_load_env()
DEEPSEEK_KEY = os.environ.get("DEEPSEEK_API_KEY", "")
DEEPSEEK_URL = "https://api.deepseek.com/chat/completions"

# ── Helpers ───────────────────────────────────────────

def detect_mode(desc):
    d = desc.lower()
    if "remote" in d:
        return "🏠 Remote" if "hybrid" not in d else "🏢🏠 Hybrid"
    return "🏢 In-office"


def call_deepseek(prompt, system="You are a helpful career coach.", max_tokens=800):
    if not DEEPSEEK_KEY:
        return None
    try:
        r = requests.post(DEEPSEEK_URL, headers={
            "Authorization": f"Bearer {DEEPSEEK_KEY}",
            "Content-Type": "application/json",
        }, json={
            "model": "deepseek-chat",
            "messages": [{"role": "system", "content": system}, {"role": "user", "content": prompt}],
            "max_tokens": max_tokens, "temperature": 0.7,
        }, timeout=45)
        if r.status_code == 200:
            return r.json()["choices"][0]["message"]["content"]
    except:
        pass
    return None


# ── Glassdoor Lookup ──────────────────────────────────

def lookup_glassdoor(company_name):
    """Try to find Glassdoor rating + salary for a company."""
    rating, salary = None, None

    # 1. Try direct Glassdoor search (public search page)
    try:
        gd_url = f"https://www.glassdoor.com/Search/results.htm?keyword={urllib.parse.quote(company_name + ' Singapore')}"
        resp = requests.get(gd_url, headers=HEADERS, timeout=12)
        if resp.status_code == 200:
            soup = BeautifulSoup(resp.text, "html.parser")
            for el in soup.find_all(["span", "div"], class_=re.compile(r"rating|score", re.I)):
                t = el.get_text().strip()
                m = re.search(r'(\d\.\d)', t)
                if m and 1 <= float(m.group(1)) <= 5:
                    rating = float(m.group(1))
                    break
            salary_text = soup.get_text()
            m = re.search(r'(?:S\$\s?|SGD\s?)([\d,]+)\s*(?:–|-|to)\s*(?:S\$\s?|SGD\s?)?([\d,]+)', salary_text, re.I)
            if m:
                salary = _to_monthly(f"SGD {m.group(1)} - {m.group(2)}")
            if rating or salary:
                return {"rating": rating, "salary": salary}
    except:
        pass

    # 2. Fallback: search engines
    for engine in ["ddg", "google", "bing"]:
        try:
            if engine == "ddg":
                url = f"https://html.duckduckgo.com/html/?q={urllib.parse.quote(company_name + ' glassdoor review')}"
            elif engine == "google":
                url = f"https://www.google.com/search?q={urllib.parse.quote(company_name + ' glassdoor review rating')}&hl=en"
            else:
                url = f"https://www.bing.com/search?q={urllib.parse.quote(company_name + ' glassdoor review')}"
            resp = requests.get(url, headers=HEADERS, timeout=10)
            if resp.status_code != 200:
                continue
            soup = BeautifulSoup(resp.text, "html.parser")

            if not rating:
                for el in soup.find_all(["span", "div", "em", "a"]):
                    t = el.get_text()
                    m = re.search(r'(\d\.\d)\s*(?:out of 5|stars|rating|★|/5)', t, re.I)
                    if m:
                        rating = float(m.group(1))
                        break

            if not salary:
                salary_text = soup.get_text()
                m = re.search(r'(?:S\$\s?|SGD\s?)([\d,]+)\s*(?:–|-|to)\s*(?:S\$\s?|SGD\s?)?([\d,]+)', salary_text, re.I)
                if m:
                    salary = _to_monthly(f"SGD {m.group(1)} - {m.group(2)}")
                else:
                    m = re.search(r'(?:S\$\s?|SGD\s?)([\d,]+)\s*(?:/yr|/year|per year|/mo|/month)', salary_text, re.I)
                    if m:
                        suffix = "/mo" if re.search(r'/mo|/month', m.group(0), re.I) else ""
                        raw = f"SGD {m.group(1)}"
                        salary = raw + suffix if suffix else _to_monthly(raw)

            if rating or salary:
                return {"rating": rating, "salary": salary}
        except:
            continue
    # 3. AI fallback: educated guess based on company + role
    return {"error": "Lookup blocked (datacenter IP rejected by search engines). Works on local runs."}


def _to_monthly(salary_str):
    """Convert an annual salary string like 'SGD 60,000 - 90,000' to monthly 'SGD 5,000 - 7,500/mo'."""
    if not salary_str:
        return salary_str
    nums = re.findall(r'([\d,]+)', salary_str)
    if not nums:
        return salary_str
    converted = []
    for n in nums:
        try:
            val = int(n.replace(',', '')) // 12
            converted.append(f"{val:,}")
        except:
            converted.append(n)
    if len(converted) >= 2:
        return f"SGD {converted[0]} - {converted[1]}/mo"
    elif converted:
        return f"SGD {converted[0]}/mo"
    return salary_str


def guess_company_info(company_name, job_title=""):
    """Use DeepSeek to estimate company rating + monthly salary when Glassdoor is unavailable."""
    if not DEEPSEEK_KEY:
        return None
    prompt = f"""Company: {company_name}
Role: {job_title}
Location: Singapore

Based on your knowledge of this company and similar roles in Singapore, estimate:
1. Company rating out of 5 (based on general reputation)
2. Estimated MONTHLY salary range in SGD for this role (e.g. "SGD 5,000 - 8,000/mo")

Format your answer exactly like this:
Rating: X.X/5
Salary: SGD X,XXX - X,XXX/mo

If you're unsure, give your best estimate and note the uncertainty."""
    result = call_deepseek(prompt, "You are a compensation analyst with knowledge of Singapore tech salaries. Always return MONTHLY salary.", max_tokens=150)
    if not result:
        return None
    info = {}
    m = re.search(r'Rating:\s*(\d\.?\d?)', result)
    if m:
        try:
            info["rating"] = float(m.group(1))
        except:
            pass
    m = re.search(r'Salary:\s*(SGD\s*[\d,]+)\s*(?:–|-|to)\s*(SGD\s*[\d,]+)', result, re.I)
    if m:
        sl = f"{m.group(1).strip()} - {m.group(2).strip()}"
        info["salary"] = sl if sl.endswith("/mo") else sl + "/mo"
    else:
        m = re.search(r'Salary:\s*(SGD\s*[\d,]+)', result, re.I)
        if m:
            sl = m.group(1).strip()
            info["salary"] = sl if sl.endswith("/mo") else sl + "/mo"
    if info:
        info["ai_guess"] = True
    return info if info else None


# ── Scraper ───────────────────────────────────────────

# (logo extraction removed)


def _parse_embedded(soup):
    """Try to extract jobs from __NEXT_DATA__, __INITIAL_STATE__, __APOLLO_STATE__, or JSON-LD."""
    jobs = []
    seen_urls = set()

    def _add_job(jp, source=""):
        """Extract a job dict from a dict that looks like a JobPosting."""
        # Safely cast everything — LinkedIn may store numbers where strings belong
        def _s(v, default=""):
            if v is None:
                return default
            if isinstance(v, str):
                return v
            if isinstance(v, (int, float, bool)):
                return str(v) if v else default
            return default

        url = ""
        entity_urn = _s(jp.get("entityUrn"))
        if entity_urn and ":" in entity_urn:
            url = f"https://www.linkedin.com/jobs/view/{entity_urn.split(':')[-1]}/"
        if not url:
            url = _s(jp.get("url")) or _s(jp.get("applyUrl"))
        if not url:
            return
        if url in seen_urls:
            return
        seen_urls.add(url)

        # ── Description: try known keys, then hunt for longest text value ──
        desc = ""
        desc_keys = ["description", "descriptionText", "jobDescription", "descriptionHtml",
                     "detailDescription", "summary", "body", "content", "text", "html"]
        for dk in desc_keys:
            raw = jp.get(dk, "")
            if isinstance(raw, str) and raw.strip():
                if raw[0] == "<" or "<" in raw[:200]:
                    desc = BeautifulSoup(raw, "html.parser").get_text().strip()
                else:
                    desc = raw.strip()
                if len(desc) > 80:
                    break
                desc = ""
            elif isinstance(raw, dict):
                for ik in ["text", "html", "plainText", "content", "body"]:
                    inner = raw.get(ik, "")
                    if isinstance(inner, str) and inner.strip():
                        if inner[0] == "<" or "<" in inner[:200]:
                            desc = BeautifulSoup(inner, "html.parser").get_text().strip()
                        else:
                            desc = inner.strip()
                        if len(desc) > 80:
                            break
                        desc = ""
                if len(desc) > 80:
                    break
        # Fallback: recursively scan ALL nested values for the longest text (>150 chars)
        if not desc:
            best = ""
            def _find_longest(obj, depth=0):
                nonlocal best
                if depth > 6:
                    return
                if isinstance(obj, str) and len(obj) > len(best) and len(obj) > 150:
                    best = obj
                elif isinstance(obj, dict):
                    for v in obj.values():
                        _find_longest(v, depth + 1)
                elif isinstance(obj, list):
                    for v in obj[:30]:
                        _find_longest(v, depth + 1)
            _find_longest(jp)
            if best:
                if best[0] == "<" or "<" in best[:200]:
                    desc = BeautifulSoup(best, "html.parser").get_text().strip()
                else:
                    desc = best.strip()

        # ── Company: try known keys, ensure it's always a string ──
        company = "Unknown"
        company_keys = ["hiringOrganization", "companyDetails", "company", "companyName",
                        "employer", "hiringCompany", "organization", "org"]
        for ck in company_keys:
            org = jp.get(ck)
            if isinstance(org, dict):
                for nk in ["name", "companyName", "title", "label"]:
                    v = org.get(nk)
                    if isinstance(v, str) and v.strip():
                        company = v.strip()
                        break
                if company != "Unknown":
                    break
            elif isinstance(org, str) and org.strip():
                company = org.strip()
                break
        if not isinstance(company, str) or company in ("0", "1", "None", "null", ""):
            company = "Unknown"

        # ── Title: try known keys ──
        title = ""
        for tk in ["title", "jobTitle", "name", "headline", "position", "role"]:
            v = jp.get(tk)
            if isinstance(v, str) and v.strip() and len(v) > 2:
                title = v.strip()
                break
        if not title:
            title = "Untitled"

        # ── Logo ── (removed)
        logo = ""

        # ── Date ──
        date_posted = ""
        for dk in ["datePosted", "listedAt", "createdAt", "publishedAt", "postDate", "postedDate", "postedAt"]:
            dv = jp.get(dk)
            if dv is not None:
                date_posted = str(dv)
                break
        if date_posted and date_posted.isdigit() and len(date_posted) >= 10:
            try:
                import datetime
                ts = int(date_posted) / 1000 if len(date_posted) > 10 else int(date_posted)
                dt = datetime.datetime.fromtimestamp(ts)
                days_ago = (datetime.datetime.now() - dt).days
                date_posted = f"{days_ago} days ago" if days_ago > 0 else "Today"
            except:
                pass

        if title == "Untitled" and not desc:
            return  # skip entries with zero useful content

        jobs.append({
            "title": title,
            "company": company,
            "url": url,
            "description": desc,
            "requirements": extract_requirements(desc),
            "mode": detect_mode(desc),
            "logo": logo,
            "date_posted": date_posted,
        })

    # 1) __NEXT_DATA__
    next_data = soup.find("script", id="__NEXT_DATA__")
    if next_data and next_data.string:
        try:
            data = json.loads(next_data.string)
            def walk(obj, depth=0):
                if depth > 20:
                    return
                if isinstance(obj, dict):
                    if "jobPosting" in obj:
                        _add_job(obj["jobPosting"])
                    # Broader: any dict with title+url in linkedin jobs context — require desc/company signal
                    keys = set(obj.keys()) if isinstance(obj, dict) else set()
                    has_job_keys = keys & {"title", "companyName", "hiringOrganization", "jobPosting", "description", "entityUrn"}
                    has_desc_signal = keys & {"description", "descriptionText", "jobDescription", "hiringOrganization", "companyName", "companyDetails"}
                    has_url = keys & {"url", "applyUrl", "entityUrn"}
                    if len(has_job_keys) >= 2 and len(has_url) >= 1 and len(has_desc_signal) >= 1:
                        url_val = str(obj.get("url", "") or obj.get("entityUrn", ""))
                        if "linkedin.com/jobs" in url_val or "linkedin.com" in url_val or "job" in url_val.lower():
                            _add_job(obj)
                    for v in obj.values():
                        walk(v, depth + 1)
                elif isinstance(obj, list):
                    for v in obj[:100]:
                        walk(v, depth + 1)
            walk(data)
        except:
            pass

    # 2) window.__INITIAL_STATE__ / window.__APOLLO_STATE__ in inline scripts
    for script in soup.find_all("script"):
        if not script.string:
            continue
        for prefix in ["window.__INITIAL_STATE__", "window.__APOLLO_STATE__"]:
            if prefix not in script.string:
                continue
            try:
                start = script.string.index(prefix)
                after = script.string[start + len(prefix):]
                # Find the JSON object
                eq = after.find("=")
                if eq == -1:
                    eq = after.find("{")
                else:
                    eq = after.find("{", eq)
                if eq == -1:
                    continue
                # extract JSON by brace matching
                depth_count = 0
                end = eq
                for i, ch in enumerate(after[eq:], eq):
                    if ch == "{":
                        depth_count += 1
                    elif ch == "}":
                        depth_count -= 1
                        if depth_count == 0:
                            end = i + 1
                            break
                if end <= eq:
                    continue
                raw = after[eq:end]
                data = json.loads(raw)
                def walk2(obj, d=0):
                    if d > 20:
                        return
                    if isinstance(obj, dict):
                        for k in ["jobs", "jobPostings", "results", "items", "data", "included"]:
                            if k in obj and isinstance(obj[k], list):
                                for item in obj[k][:50]:
                                    if isinstance(item, dict):
                                        _add_job(item)
                        # Walk all dict values
                        for v in obj.values():
                            walk2(v, d + 1)
                    elif isinstance(obj, list):
                        for v in obj[:100]:
                            walk2(v, d + 1)
                walk2(data)
            except:
                continue
    # 3) JSON-LD
    for s in soup.find_all("script", type="application/ld+json"):
        try:
            data = json.loads(s.string)
            items = data if isinstance(data, list) else [data]
            for item in items:
                if item.get("@type") == "JobPosting":
                    _add_job(item)
        except:
            pass

    # 4) Hunt raw page text for additional jobs (works when structured data is sparse)
    page_text = soup.get_text()
    # Look for job title patterns followed by description-like text
    # LinkedIn pages often have job titles in h1/h2 with desc following
    for h in soup.find_all(["h1", "h2", "h3"]):
            title = h.get_text(strip=True)
            if not title or len(title) < 3 or len(title) > 120:
                continue
            # Look for a substantial text block after this heading
            parent = h.find_parent(["div", "section", "article"])
            if not parent:
                continue
            desc_candidates = parent.find_all(["p", "div", "span", "section", "article"])
            for dc in desc_candidates:
                txt = dc.get_text(" ", strip=True)
                if len(txt) > 200:
                    # Try to find a URL near this heading
                    link = parent.find("a", href=re.compile(r"/jobs/"))
                    url = ""
                    if link:
                        url = link["href"]
                        if not url.startswith("http"):
                            url = "https://www.linkedin.com" + url
                    jobs.append({
                        "title": title,
                        "company": "Unknown",
                        "url": url,
                        "description": txt[:3000],
                        "requirements": extract_requirements(txt),
                        "mode": detect_mode(txt),
                        "logo": "",
                        "date_posted": "",
                    })
                    break
            if len(jobs) >= 10:
                break
def _extract_card_date(card):
    """Extract posting date / time-ago text from a job card."""
    for el in card.find_all(["time", "span", "p"]):
        text = el.get_text(strip=True).lower()
        if any(w in text for w in ["day", "week", "month", "hour", "minute", "ago", "just now"]):
            return el.get_text(strip=True)
        if el.get("datetime"):
            return el["datetime"]
    return ""


def _extract_card_snippet(card):
    """Extract visible description snippet from a job card. Returns '' if only location/date found."""
    def _is_metadata(text):
        """Reject text that looks like location, date, or time-ago."""
        t = text.lower().strip()
        # Pure location patterns
        if re.match(r'^[a-z]+(,\s*[a-z]+)+$', t) and len(t) < 60:
            return True
        # Date / time-ago patterns
        if re.search(r'\d+\s*(day|week|month|hour|minute|yr|year)s?\s*ago|today|yesterday|just now', t):
            if len(t) < 80:
                return True
        # Pure date like "2026-05-13"
        if re.match(r'^\d{4}-\d{2}-\d{2}$', t.strip()):
            return True
        # Singapore, Singapore (duplicate city/country)
        if re.match(r'^[A-Z][a-z]+,\s*[A-Z][a-z]+(\s+\d+)?$', t.strip()) and len(t) < 40:
            return True
        return False

    # 1) Try class-based matching
    for cls_pat in [r"snippet", r"metadata", r"description", r"summary", r"detail", r"info", r"body", r"text"]:
        el = card.find(["p", "div", "span", "section"], class_=re.compile(cls_pat, re.I))
        if el:
            text = el.get_text(" ", strip=True)
            if len(text) > 80 and not _is_metadata(text):
                return text
    # 2) Try any <p> or <span> with substantial text
    for tag in ["p", "span", "div"]:
        for el in card.find_all(tag):
            text = el.get_text(" ", strip=True)
            if len(text) > 80:
                if not re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+){0,3}$', text):
                    if not _is_metadata(text):
                        return text
    # 3) Get all text from the card, strip known short lines (title, company, location, date)
    all_text = card.get_text("\n", strip=True)
    lines = [l.strip() for l in all_text.split("\n") if l.strip()]
    meaningful = [l for l in lines if len(l) > 60
                  and not re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+){0,3}$', l)
                  and not _is_metadata(l)]
    if meaningful:
        return " ".join(meaningful[:5])
    # 4) Try parent elements
    parent = card.find_parent(["div", "li", "section"])
    if parent:
        for el in parent.find_all(["p", "span", "div"]):
            text = el.get_text(" ", strip=True)
            if len(text) > 100:
                if not re.match(r'^[A-Z][a-z]+(\s+[A-Z][a-z]+){0,3}$', text):
                    if not _is_metadata(text):
                        return text[:2000]
    return ""


@st.cache_data(ttl=600, show_spinner=False)
def scrape_linkedin(keywords, max_jobs=30):
    jobs = []
    seen = set()
    for start in range(0, max_jobs, 25):
        url = (
            f"https://www.linkedin.com/jobs/search/?"
            f"keywords={urllib.parse.quote(keywords)}"
            f"&location={urllib.parse.quote(LOCATION)}"
            f"&f_TPR={TIME_RANGE}"
            f"&position=1&pageNum={start // 25}"
        )
        page_jobs = 0
        for headers in (HEADERS, GOOGLEBOT_HEADERS):
            if page_jobs >= max(5, max_jobs // 4):
                break
            try:
                resp = requests.get(url, headers=headers, timeout=15)
                if resp.status_code != 200:
                    continue
                soup = BeautifulSoup(resp.text, "html.parser")
            except:
                continue

            # 1) Try embedded data — no extra HTTP requests needed
            embedded = _parse_embedded(soup)
            for j in (embedded or []):
                u = j.get("url", "")
                if u and u not in seen:
                    seen.add(u)
                    jobs.append(j)
            if len(jobs) >= max_jobs:
                break

            # 2) Parse ALL visible cards for jobs the embedded data missed
            cards = soup.find_all("div", class_=re.compile(r"base-card|job-card|job-search-card"))
            if not cards:
                cards = soup.find_all("li", class_=re.compile(r"job|result"))
            for card in cards:
                if len(jobs) >= max_jobs:
                    break
                title_el = card.find(["h3", "a"], class_=re.compile(r"title|job", re.I))
                company_el = card.find(["h4", "span"], class_=re.compile(r"company|subtitle|employer", re.I))
                link_el = card.find("a", href=re.compile(r"/jobs/view/|/jobs/"))
                if not title_el or not link_el:
                    continue
                job_url = link_el["href"]
                if not job_url.startswith("http"):
                    job_url = "https://www.linkedin.com" + job_url
                job_url = job_url.split("?")[0]
                if job_url in seen:
                    continue
                seen.add(job_url)
                company = company_el.get_text(strip=True) if company_el else "Unknown"
                date_posted = _extract_card_date(card)
                desc, reqs = get_job_details(job_url)
                if not desc:
                    desc = _extract_card_snippet(card)
                # If still no description, generate a placeholder via AI
                if not desc and DEEPSEEK_KEY:
                    title = title_el.get_text(strip=True)
                    desc = call_deepseek(
                        f"Write a 2-3 sentence description of the role '{title}' at {company} in Singapore. "
                        "Describe typical responsibilities based on the job title. Be specific and professional.",
                        "You describe job roles concisely.", max_tokens=250
                    ) or ""
                    if desc:
                        reqs = extract_requirements(desc)
                jobs.append({
                    "title": title_el.get_text(strip=True),
                    "company": company,
                    "url": job_url, "description": desc,
                    "requirements": reqs, "mode": detect_mode(desc),
                    "logo": "", "date_posted": date_posted,
                })
            page_jobs = len(jobs)
            if len(jobs) >= max_jobs:
                break
            time.sleep(1.5)
        if len(jobs) >= max_jobs:
            break
    return jobs


def _clean_html(html_str):
    """Unescape LinkedIn's double-encoded HTML, keeping formatting tags for rich display."""
    import html as _html
    text = _html.unescape(html_str)
    # Remove scripts/styles/event handlers (safety)
    text = re.sub(r'<script[^>]*>.*?</script>', '', text, flags=re.I | re.DOTALL)
    text = re.sub(r'<style[^>]*>.*?</style>', '', text, flags=re.I | re.DOTALL)
    text = re.sub(r'\s+on\w+="[^"]*"', '', text, flags=re.I)
    # Keep formatting tags: br, strong, b, em, i, u, ul, ol, li, p, h1-h6
    # Strip all other tags
    text = re.sub(r'<(?!br\b|strong\b|b\b|em\b|i\b|u\b|ul\b|ol\b|li\b|p\b|h[1-6]\b|/\s*(?:br|strong|b|em|i|u|ul|ol|li|p|h[1-6])\s*)[^>]+>', '', text, flags=re.I)
    return text.strip()


def get_job_details(job_url):
    """Fetch full job description from individual LinkedIn job page."""
    # Try Googlebot UA first (LinkedIn allows search crawlers)
    for ua_headers in [GOOGLEBOT_HEADERS, HEADERS]:
        try:
            resp = requests.get(job_url, headers=ua_headers, timeout=15, allow_redirects=True)
            if resp.status_code != 200:
                continue
            if len(resp.text) < 500 or "authwall" in resp.text.lower():
                continue
            soup = BeautifulSoup(resp.text, "html.parser")

            # 1) Try visible description divs
            desc_el = (soup.find("div", class_="description__text") or
                       soup.find("div", class_="show-more-less-html__markup") or
                       soup.find("div", class_=re.compile(r"description", re.I)) or
                       soup.find("div", class_=re.compile(r"jobs-description", re.I)) or
                       soup.find("section", class_=re.compile(r"description", re.I)))
            if desc_el:
                desc = desc_el.get_text("\n", strip=True)
                if len(desc) > 200:
                    return desc, extract_requirements(desc)

            # 2) Try JSON-LD
            for s in soup.find_all("script", type="application/ld+json"):
                try:
                    data = json.loads(s.string)
                    if "description" in data:
                        desc = data["description"]
                        if isinstance(desc, str) and len(desc) > 200:
                            desc = _clean_html(desc)
                            return desc, extract_requirements(desc)
                except:
                    pass

            # 3) Try meta description
            meta = soup.find("meta", attrs={"name": "description"})
            if meta and meta.get("content"):
                desc = meta["content"].strip()
                if len(desc) > 200:
                    return desc, extract_requirements(desc)

            # 4) Try inline script with job data
            for script in soup.find_all("script"):
                if not script.string:
                    continue
                if '"description"' in script.string and '"title"' in script.string:
                    try:
                        # Find JSON object containing description
                        for m in re.finditer(r'\{[^{}]*"description"\s*:\s*"((?:[^"\\]|\\.)*)"[^}]*\}', script.string):
                            raw_desc = m.group(1)
                            if len(raw_desc) > 200:
                                desc = _clean_html(raw_desc)
                                return desc, extract_requirements(desc)
                    except:
                        pass
        except:
            continue
    return "", ""


def extract_requirements(text):
    m = re.search(r'(?:requirement|qualification|what you.{0,15}need|must have|required.{0,10}skill)', text, re.IGNORECASE)
    if m:
        section = text[m.start():]
        stop = len(section)
        for kw in ['responsibilities', 'about the role', 'we offer', 'benefits']:
            m2 = re.search(kw, section[80:], re.IGNORECASE)
            if m2:
                stop = min(stop, 80 + m2.start())
        return section[:stop].strip()
    lines = text.split('\n')
    req_lines = [l for l in lines if any(w in l.lower() for w in ['require', 'qualif', 'skill', 'experience', 'degree', 'year'])]
    return '\n'.join(req_lines[:10]) if req_lines else text[:400]


# ── Resume Parse ──────────────────────────────────────

def parse_resume(uploaded_file):
    ext = os.path.splitext(uploaded_file.name)[1].lower()
    if ext == '.pdf':
        from PyPDF2 import PdfReader
        reader = PdfReader(uploaded_file)
        return "\n".join((page.extract_text() or "") for page in reader.pages).strip()
    if ext == '.docx':
        doc = Document(uploaded_file)
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip()).strip()
    if ext == '.doc':
        return _try_read_doc(uploaded_file)
    return uploaded_file.read().decode("utf-8", errors="ignore").strip()


def _try_read_doc(uploaded_file):
    content = uploaded_file.read()
    text = content.decode("utf-8", errors="ignore")
    lines = [l.strip() for l in text.split('\n') if len(l.strip()) > 3 and not all(c < ' ' or c > '~' or c == '\x00' for c in l.strip())]
    result = '\n'.join(lines)
    if len(result) > 50:
        return result
    try:
        import subprocess
        with tempfile.NamedTemporaryFile(suffix='.doc', delete=False) as tf:
            tf.write(content)
            tf.flush()
            out = subprocess.check_output(['antiword', tf.name], timeout=10)
            os.unlink(tf.name)
            return out.decode("utf-8", errors="ignore").strip()
    except:
        pass
    return result if len(result) > 50 else ""


# ── Resume Tailoring ──────────────────────────────────

def tailor_resume_simple(resume, job):
    job_text = f"{job['title']} {job['description']} {job['requirements']}".lower()
    stop_words = {'this','that','with','from','your','have','will','they','about','their','would','which','the','and','for','are','but','not','you','all','can','had','has','was','were','been','being','more','some','than','then','its','also','what','when','where','who','how','our'}
    job_words = set(re.findall(r'\b[a-z]{4,}\b', job_text)) - stop_words
    resume_words = set(re.findall(r'\b[a-z]{4,}\b', resume.lower()))
    missing = sorted(job_words - resume_words)
    rate = len(job_words & resume_words) / max(len(job_words), 1) * 100
    return rate, missing[:20]


def tailor_resume_ai(resume, job):
    prompt = f"""Job: {job['title']} at {job['company']}
Description: {job.get('description','')[:2000]}
Requirements: {job.get('requirements','')[:1000]}

My current resume:
{resume[:3000]}

Analyze my resume against this job. Return:
1. Match percentage (estimate)
2. Top 5 missing skills/keywords I should add
3. 3 bullet points to add to my resume
4. Brief summary of my strengths
Keep it concise."""
    return call_deepseek(prompt, "You are an expert resume reviewer. Be honest and specific.")


def generate_tailored_resume(resume, job):
    # Build section list from user's checkbox preferences (default: no certs)
    sections = ["NAME: [Full Name from resume]"]
    if st.session_state.get("sec_summary", True):
        sections.append("SUMMARY: [2-3 sentence professional summary tailored to this role]")
    if st.session_state.get("sec_skills", True):
        sections.append("SKILLS: [comma-separated skills, prioritize those matching the job]")
    if st.session_state.get("sec_experience", True):
        sections.append("EXPERIENCE: [keep same roles, rewrite bullets to match job keywords]")
    if st.session_state.get("sec_education", True):
        sections.append("EDUCATION: [keep as-is]")
    if st.session_state.get("sec_certs", False):
        sections.append("CERTIFICATIONS: [keep as-is, if any]")
    
    prompt = f"""Job Title: {job['title']}
Company: {job['company']}
Job Description: {job.get('description','')[:2000]}
Requirements: {job.get('requirements','')[:1000]}

My current resume:
{resume[:3000]}

Rewrite my resume for this job. Rules:
- NEVER add anything I didn't actually do — no invented skills, metrics, or achievements
- Keep my real job titles, dates, and company names exactly as-is
- Use professional language with strong action verbs (e.g. "Led", "Designed", "Built", "Optimised")
- Use UK English spelling throughout: -ise not -ize, -our not -or, -re not -er, "travelling" not "traveling"
- Avoid semicolons and run-on sentences — keep each bullet to one clear idea
- Weave ATS keywords from the job description naturally into the summary and bullets
- Use the employer's exact terminology for technical terms (e.g. "CI/CD pipelines" not "deployment automation"), but apply UK spelling to all other words
Include ONLY these sections in this exact order:
- NEVER add anything I didn't actually do — no invented skills, metrics, or achievements
- Keep my real job titles, dates, and company names exactly as-is
- Use plain, direct language — short sentences, no semicolons, no buzzwords like "leveraged" or "spearheaded"
- Write like a human writes: "I built the API" not "Spearheaded the development of a RESTful API"
- Keep bullets concise — one line each when possible
- Weave in ATS keywords naturally from the job description
Include ONLY these sections in this exact order:

{chr(10).join(sections)}

IMPORTANT — ATS Optimisation: Include keywords and phrases from the job description naturally in the SUMMARY and SKILLS sections. Use exact terminology the employer uses (e.g. if they say "CI/CD pipelines", use that phrase, not "deployment automation"). Place the most relevant keywords early in each section. Return the full rewritten resume with those exact headers."""
    return call_deepseek(prompt, "You are an expert resume writer. Use UK English spelling throughout (-ise not -ize, -our not -or, -re not -er). Use strong professional action verbs. Avoid semicolons and run-on sentences. Weave ATS keywords naturally. NEVER add anything not in the original resume.", max_tokens=1500)


def build_docx(tailored_text):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    section_headers = {"SUMMARY:", "SKILLS:", "EXPERIENCE:", "EDUCATION:", "CERTIFICATIONS:"}
    lines = tailored_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("NAME:"):
            p = doc.add_paragraph()
            run = p.add_run(line.replace("NAME:", "").strip())
            run.bold = True
            run.font.size = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        elif any(line.startswith(h) for h in section_headers):
            p = doc.add_paragraph()
            parts = line.split(":", 1)
            run = p.add_run(parts[0] + ":")
            run.bold = True
            run.font.size = Pt(13)
            if len(parts) > 1 and parts[1].strip():
                p.add_run("  " + parts[1].strip())

        elif line.startswith("- ") or line.startswith("\u2022 "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            p.paragraph_format.space_after = Pt(2)
            text = line.lstrip("-\u2022 ").strip()
            p.add_run("\u2022 " + text)

        else:
            p = doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf





# ── Interview ─────────────────────────────────────────

def generate_questions_ai(job):
    prompt = f"""Job: {job['title']} at {job['company']}
Description: {job.get('description','')[:2000]}
Requirements: {job.get('requirements','')[:1000]}

Generate 8 interview questions. Prioritize behavioral and situational (5-6), with only 2-3 technical. Make them specific to this role and company. Return as numbered list. Do NOT repeat the question number in the question text."""
    result = call_deepseek(prompt, "You are a senior hiring manager focused on behavioral and culture-fit assessment. Be specific and challenging.")
    if result:
        qs = []
        for line in result.split('\n'):
            line = line.strip()
            if not line:
                continue
            cleaned = re.sub(r'^\d+[\.\)\-]+\s*', '', line)
            if cleaned and len(cleaned) > 10:
                qs.append(cleaned)
        return qs[:10]
    return None


def generate_questions_simple(job):
    desc = (job.get('description', '') + ' ' + job.get('requirements', '')).lower()
    title = job.get('title', '').lower()
    qs = []
    # 1-2 light tech questions phrased as behavioral
    tech_map = {
        'python': "Walk me through how you'd debug a performance issue in a Python service.",
        'javascript': "Describe a time you had to optimize a slow web application.",
        'java': "Tell me about a complex system you built in Java and the trade-offs you made.",
        'sql': "Describe how you'd investigate a slow database query in production.",
        'aws': "Tell me about a time you designed or migrated a cloud architecture.",
        'docker': "Walk me through how you set up CI/CD for a containerized app.",
        'react': "Describe a challenging UI problem you solved and your approach.",
        'data': "Tell me about a time your data analysis changed a business decision.",
        'api': "Describe how you designed an API and handled versioning or breaking changes.",
        'cloud': "Tell me about a system you scaled and the bottlenecks you hit.",
    }
    for kw, q in tech_map.items():
        if kw in desc and len(qs) < 2:
            qs.append(q)
    # Behavioral questions (majority)
    behavioral = [
        "Tell me about yourself and your background — focus on what brought you to this field.",
        "Why are you interested in this role and this company specifically?",
        "Describe the most challenging project you worked on. What made it hard and how did you overcome it?",
        "Tell me about a time you disagreed with a coworker or manager. How did you handle it?",
        "Describe a situation where you had to learn something completely new on the job.",
        "Tell me about a time you failed or made a mistake. What happened and what did you learn?",
        "How do you prioritize when you have multiple competing deadlines?",
        "Describe a time you had to influence a team or stakeholder without formal authority.",
    ]
    qs.extend(behavioral)
    if 'senior' in title or 'lead' in title:
        qs.extend([
            "Tell me about a time you had to manage an underperforming team member.",
            "How do you balance hands-on technical work with leadership responsibilities?",
        ])
    if 'manager' in title or 'lead' in title:
        qs.append("Describe your approach to building and scaling a high-performing team.")
    return qs[:10]


def evaluate_answer(question, answer):
    if not DEEPSEEK_KEY or len(answer.split()) < 5:
        return None
    prompt = f"""Question: {question}
Candidate's answer: {answer}
Brief feedback (2-3 sentences): what was strong, what to improve, score 1-10."""
    return call_deepseek(prompt, "You are an interview coach. Be constructive and brief.")


# ── UI ────────────────────────────────────────────────


def _extract_title(text, use_ai=True):
    """Extract job title from pasted job description or LinkedIn URL."""
    if not text:
        return "Custom Job"
    # Try LinkedIn URL pattern
    m = re.search(r'linkedin\.com/jobs/view/([^/]+)-\d+', text)
    if m:
        return m.group(1).replace('-', ' ').title()
    # Try AI extraction
    if use_ai and DEEPSEEK_KEY:
        result = call_deepseek(
            f"Extract ONLY the job title from this text. Return just the title, nothing else:\n\n{text[:1500]}",
            "You extract job titles. Return only the title, no explanation.", max_tokens=50
        )
        if result and len(result.strip()) > 2 and len(result.strip()) < 80:
            return result.strip()
    # Fallback: first substantial line that isn't metadata
    lines = [l.strip() for l in text.split('\n') if l.strip() and len(l.strip()) > 3]
    for line in lines[:5]:
        if re.match(r'^(at |@|[A-Z][a-z]+(?: [A-Z][a-z]+){0,2}$)', line):
            continue
        if re.match(r'^(Singapore|Remote|Hybrid|In.office|Full.time|Part.time|Contract|S\$|Posted|\d+ (day|week|month)s? ago)', line, re.I):
            continue
        if 5 < len(line) < 80 and not line.startswith('http'):
            return line[:80]
    return "Custom Job"


def _extract_company(text, use_ai=True):
    """Extract company name from pasted job description."""
    if not text:
        return "Unknown"
    # Try AI extraction
    if use_ai and DEEPSEEK_KEY:
        result = call_deepseek(
            f"Extract ONLY the company name from this job posting. Return just the company name, nothing else:\n\n{text[:1500]}",
            "You extract company names. Return only the name, no explanation.", max_tokens=30
        )
        if result and len(result.strip()) > 1 and len(result.strip()) < 50:
            return result.strip()
    # Fallback regex
    m = re.search(r'\bat\s+([A-Z][A-Za-z0-9 .&,-]+?)(?:\s+is\b|\s+in\b|\s+Singapore|\s*$|\.)', text)
    if m and len(m.group(1)) > 2:
        return m.group(1).strip().rstrip('.,')
    m = re.search(r'@\s*([A-Z][A-Za-z0-9 .&,-]+?)(?:\s|$|\.|,)', text)
    if m and len(m.group(1)) > 2:
        return m.group(1).strip()
    m = re.search(r'([A-Z][A-Za-z0-9 .&,-]{2,30})\s+is\s+(hiring|looking)', text)
    if m:
        return m.group(1).strip()
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:5]:
        if re.match(r'^[A-Z][A-Za-z0-9 .&,-]{2,30}$', line) and not re.match(r'^(Singapore|Remote|Hybrid)', line, re.I):
            return line[:40]
    return "Unknown"

    if m and len(m.group(1)) > 2:
        return m.group(1).strip().rstrip('.,')
    m = re.search(r'@\s*([A-Z][A-Za-z0-9 .&,-]+?)(?:\s|$|\.|,)', text)
    if m and len(m.group(1)) > 2:
        return m.group(1).strip()
    m = re.search(r'([A-Z][A-Za-z0-9 .&,-]{2,30})\s+is\s+(hiring|looking)', text)
    if m:
        return m.group(1).strip()
    # Try first line that looks like a company name
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    for line in lines[:5]:
        if re.match(r'^[A-Z][A-Za-z0-9 .&,-]{2,30}$', line) and not re.match(r'^(Singapore|Remote|Hybrid)', line, re.I):
            return line[:40]
    return "Unknown"


def build_tailored_docx(original_bytes, tailored_text):
    """Edit the original DOCX file in-place with tailored content, preserving formatting."""
    try:
        from docx import Document as DocxDocument
        from io import BytesIO
        
        doc = DocxDocument(BytesIO(original_bytes))
        # Parse tailored text into sections
        sections = {}
        current_section = "OTHER"
        for line in tailored_text.split('\n'):
            line = line.strip()
            if not line:
                continue
            if line.startswith("NAME:"):
                sections["NAME"] = line.replace("NAME:", "").strip()
            elif any(line.startswith(h) for h in ["SUMMARY:", "SKILLS:", "EXPERIENCE:", "EDUCATION:", "CERTIFICATIONS:"]):
                parts = line.split(":", 1)
                current_section = parts[0]
                sections[current_section] = parts[1].strip() if len(parts) > 1 else ""
            else:
                sections[current_section] = sections.get(current_section, "") + "\n" + line
        
        # Walk through paragraphs and replace content by section
        section_order = ["NAME", "SUMMARY", "SKILLS", "EXPERIENCE", "EDUCATION", "CERTIFICATIONS"]
        section_idx = 0
        found_name = False
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # Detect section by bold/header patterns
            upper_text = text.upper()
            for sec in section_order:
                if upper_text.startswith(sec + ":") or upper_text == sec:
                    if sec in sections and sections[sec]:
                        # Replace the section header + content
                        new_text = f"{sec}: {sections[sec]}"
                        # Clear paragraph and set new text, preserving formatting
                        for run in para.runs:
                            run.text = ""
                        if para.runs:
                            para.runs[0].text = new_text
                        else:
                            para.add_run(new_text)
                    section_idx = section_order.index(sec) + 1
                    found_name = (sec == "NAME")
                    break
            else:
                # Body text under current section — append to first non-header paragraph
                if section_idx > 0 and not found_name:
                    pass  # Skip body text replacement for now (complex)
        
        buf = BytesIO()
        doc.save(buf)
        buf.seek(0)
        return buf
    except Exception:
        return None


def scrape_similar_jobs(job_title, company, max_jobs=5):
    """Find similar jobs on LinkedIn by searching with the same title."""
    query = urllib.parse.quote(f"{job_title} -{company}")
    url = f"https://www.linkedin.com/jobs/search/?keywords={query}&location=Singapore&f_TPR=r1209600&position=1&pageNum=0"
    try:
        resp = requests.get(url, headers=GOOGLEBOT_HEADERS, timeout=15)
        if resp.status_code != 200:
            return []
        soup = BeautifulSoup(resp.text, "html.parser")
        cards = soup.find_all("div", class_=re.compile(r"base-card|job-card"))
        if not cards:
            cards = soup.find_all("li", class_=re.compile(r"job|result"))
        
        similar = []
        seen_titles = set()
        for card in cards:
            if len(similar) >= max_jobs:
                break
            title_el = card.find(["h3", "h2", "span"], class_=re.compile(r"title", re.I))
            company_el = card.find(["h4", "span", "div"], class_=re.compile(r"subtitle|company", re.I))
            link_el = card.find("a", href=re.compile(r"/jobs/"))
            if not title_el:
                continue
            title = title_el.get_text(strip=True)
            comp = company_el.get_text(strip=True) if company_el else "?"
            if title.lower() in seen_titles or title.lower() == job_title.lower():
                continue
            seen_titles.add(title.lower())
            url = ""
            if link_el:
                url = link_el["href"]
                if not url.startswith("http"):
                    url = "https://www.linkedin.com" + url
            similar.append({"title": title, "company": comp, "url": url})
        return similar
    except:
        return []

def main():
    st.title("🎯 Job Scraper + Resume Tailor + Mock Interview")
    st.caption("Scrapes LinkedIn jobs in Singapore (last 2 weeks) | Powered by DeepSeek")

    # Keyboard navigation — left/right arrow keys to browse jobs
    st.markdown("""
    <script>
    (function() {
        function clickBtn(text) {
            var btns = window.parent.document.querySelectorAll('button');
            for (var i = 0; i < btns.length; i++) {
                if (btns[i].innerText && btns[i].innerText.indexOf(text) !== -1) {
                    btns[i].click();
                    return true;
                }
            }
            return false;
        }
        document.addEventListener('keydown', function(e) {
            if (e.target.tagName === 'INPUT' || e.target.tagName === 'TEXTAREA') return;
            if (e.key === 'ArrowLeft') { clickBtn('\u25c0'); e.preventDefault(); }
            if (e.key === 'ArrowRight') { clickBtn('\u25b6'); e.preventDefault(); }
        });
    })();
    </script>
    """, unsafe_allow_html=True)

    # Sidebar
    with st.sidebar:
        st.header("⚙️ Settings")
        kw = st.text_input("Job title / keywords", placeholder="e.g. software engineer")
        company_filter = st.text_input("Company (optional)", placeholder="e.g. Google, Shopee")
        max_jobs = st.slider("Max jobs", 10, 50, 25)
        st.divider()

        st.subheader("📋 Custom Job")
        with st.expander("Paste a job for analysis", expanded=False):
            custom_url = st.text_input("LinkedIn job URL", placeholder="https://www.linkedin.com/jobs/view/...", key="custom_url")
            custom_desc = st.text_area("Or paste job description", placeholder="Paste the full job posting here...", height=120, key="custom_desc")
            if st.button("🔍 Analyze This Job", use_container_width=True, key="custom_btn"):
                if custom_url:
                    with st.spinner("Fetching job..."):
                        desc, reqs = get_job_details(custom_url)
                        if desc:
                            st.session_state.custom_job = {"title": _extract_title(desc), "company": _extract_company(desc), "url": custom_url, "description": desc, "requirements": reqs, "mode": detect_mode(desc), "logo": "", "date_posted": "", "source": "custom"}
                            st.session_state.jobs = [st.session_state.custom_job]
                            st.session_state.job_idx = 0
                            st.session_state.glassdoor_cache = {}
                            st.rerun()
                        else:
                            st.error("Could not fetch job — site may block datacenter IPs.")
                elif custom_desc.strip():
                    st.session_state.custom_job = {"title": _extract_title(custom_desc), "company": _extract_company(custom_desc), "url": "", "description": custom_desc, "requirements": extract_requirements(custom_desc), "mode": detect_mode(custom_desc), "logo": "", "date_posted": "", "source": "custom"}
                    st.session_state.jobs = [st.session_state.custom_job]
                    st.session_state.job_idx = 0
                    st.session_state.glassdoor_cache = {}
                    st.rerun()

        st.divider()
        st.subheader("📄 Resume")
        resume_file = st.file_uploader("Upload (PDF/DOCX/DOC/TXT)", type=["pdf", "docx", "doc", "txt", "md"])
        if resume_file:
            if "resume_filename" not in st.session_state or st.session_state.resume_filename != resume_file.name:
                st.session_state.resume_filename = resume_file.name
                st.session_state.resume_bytes = resume_file.getvalue()
                st.session_state.resume_text = parse_resume(resume_file)
                st.session_state.tailored_text = None
                if st.session_state.resume_text:
                    st.success(f"Loaded: {resume_file.name} ({len(st.session_state.resume_text)} chars)")
                else:
                    st.error("Could not read file. For .doc, try saving as .docx first.")

        # AI always on when connected
        if DEEPSEEK_KEY:
            st.success("DeepSeek AI connected")
            st.session_state.use_ai_tailor = True
            st.session_state.use_ai_questions = True
        else:
            st.warning("Add DEEPSEEK_API_KEY to .env or st.secrets then restart")
            st.session_state.use_ai_tailor = False
            st.session_state.use_ai_questions = False

    if not kw and not company_filter.strip() and not st.session_state.get("custom_job"):
        st.info("👈 Enter a job title or company in the sidebar to start")
        return

    # Build search query with company filter
    search_query = kw
    if company_filter.strip():
        search_query = f"{kw} {company_filter.strip()}" if kw else company_filter.strip()

    if kw or company_filter.strip():
        if st.button("🔍 Find Jobs", type="primary", use_container_width=True):
            st.session_state.jobs = None
            st.session_state.job_idx = 0
            st.session_state.glassdoor_cache = {}
            st.session_state.custom_job = None
            with st.spinner(f"Searching for '{search_query}'..."):
                st.session_state.jobs = scrape_linkedin(search_query, max_jobs)
            st.rerun()

    if "jobs" not in st.session_state:
        return
    if not st.session_state.jobs:
        st.warning("No jobs found — try different keywords or a broader search.")
        return

    jobs = st.session_state.jobs
    if "job_idx" not in st.session_state:
        st.session_state.job_idx = 0  # show first job immediately

    idx = st.session_state.job_idx
    total = len(jobs)
    # ── Job List (compact) ──
    st.success(f"Found {total} jobs")

    # Quick-jump dropdown
    job_labels = []
    for i, j in enumerate(jobs):
        meta = []
        if j.get("company"):
            meta.append(j["company"])
        if j.get("date_posted"):
            meta.append(j["date_posted"])
        label = f"{i+1}. {j['title']}  —  {', '.join(meta)}" if meta else f"{i+1}. {j['title']}"
        job_labels.append(label)
    selected = st.selectbox("Jump to job", options=range(total), format_func=lambda i: job_labels[i], index=idx, label_visibility="collapsed")
    if selected != idx:
        st.session_state.job_idx = selected
        st.rerun()

    job = jobs[idx]

    # ── Anchor for scroll-to-top on navigation ──
    st.markdown('<div id="job-top"></div>', unsafe_allow_html=True)
    st.markdown("""
    <script>
    document.getElementById('job-top')?.scrollIntoView({behavior: 'smooth'});
    </script>
    """, unsafe_allow_html=True)

    # Prev / Next
    st.divider()
    nav_cols = st.columns([1, 1])
    with nav_cols[0]:
        if st.button("◀  Prev", use_container_width=True, key="prev_btn", disabled=(idx == 0)):
            st.session_state.job_idx = max(0, idx - 1)
            st.rerun()
    with nav_cols[1]:
        if st.button("Next  ▶", use_container_width=True, key="next_btn", disabled=(idx >= total - 1)):
            st.session_state.job_idx = min(total - 1, idx + 1)
            st.rerun()

    # ── Job Card ──
    st.divider()
    info_col, link_col = st.columns([6, 1.4])
    with info_col:
        st.markdown(f"## {job['title']}")
        card_meta = [f"**{job['company']}**", job.get('mode', '')]
        if job.get("date_posted"):
            card_meta.append(f"🕒 {job['date_posted']}")
        src = job.get("source", "")
        if src:
            card_meta.append(f"via {src.title()}")
        st.markdown("  ·  ".join(p for p in card_meta if p))
    with link_col:
        if job.get("url"):
            st.link_button("🔗 Open on LinkedIn", job['url'])

    # ── Tabs ──
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["📋 Details", "📄 Tailor Resume", "❓ Questions", "🎤 Mock Interview", "💬 Ask Them"])

    with tab1:
        st.subheader("Description")
        desc_text = job.get('description') or '*No description*'
        st.markdown(desc_text[:5000], unsafe_allow_html=True)
        if job.get('requirements'):
            st.subheader("Requirements")
            st.markdown(job['requirements'][:2000])


            # ── Similar Jobs ──
            if "similar_jobs" not in st.session_state:
                st.session_state.similar_jobs = {}
            job_key = job.get("url", "") + job.get("title", "")
            if job_key not in st.session_state.similar_jobs:
                with st.spinner("Finding similar jobs..."):
                    st.session_state.similar_jobs[job_key] = scrape_similar_jobs(job["title"], job["company"])
            similar = st.session_state.similar_jobs.get(job_key, [])
            if similar:
                with st.expander(f"🔗 Similar Jobs on LinkedIn ({len(similar)})", expanded=False):
                    for sj in similar:
                        if sj.get("url"):
                            st.markdown(f"- **{sj['title']}** — *{sj['company']}*  [View]({sj['url']})")
                        else:
                            st.markdown(f"- **{sj['title']}** — *{sj['company']}*")
        # ── Glassdoor Panel ──
        with st.expander("🏢 Company Info", expanded=False):
            if "glassdoor_cache" not in st.session_state:
                st.session_state.glassdoor_cache = {}
            company = job['company']
            if company not in st.session_state.glassdoor_cache:
                with st.spinner(f"Looking up {company} on Glassdoor..."):
                    gd = lookup_glassdoor(company)
                    if gd and gd.get("error") and DEEPSEEK_KEY:
                        # Glassdoor blocked — try AI guess
                        ai = guess_company_info(company, job.get("title", ""))
                        if ai:
                            gd = ai
                    st.session_state.glassdoor_cache[company] = gd
            gd = st.session_state.glassdoor_cache.get(company)
            if gd and gd.get("error"):
                st.caption(gd["error"])
            elif gd:
                c1, c2 = st.columns(2)
                with c1:
                    if gd.get("rating"):
                        stars = "★" * int(gd["rating"]) + "☆" * (5 - int(gd["rating"]))
                        label = "AI Est. Rating" if gd.get("ai_guess") else "Glassdoor Rating"
                        st.metric(label, f"{gd['rating']:.1f} / 5")
                        st.caption(stars)
                    else:
                        st.caption("No rating found")
                with c2:
                    if gd.get("salary"):
                        label = "AI Est. Monthly" if gd.get("ai_guess") else "Est. Monthly Salary"
                        st.metric(label, gd["salary"])
                    else:
                        st.caption("No salary data")
                if gd.get("ai_guess"):
                    st.caption("⚠️ AI-generated estimate — not from Glassdoor")
            else:
                st.caption("No company data found")

    with tab2:
        st.subheader("Resume Tailoring")
        if "resume_text" not in st.session_state or not st.session_state.resume_text:
            st.warning("Upload your resume in the sidebar first")
        else:
            st.caption("Select sections to include:")
            sc1, sc2, sc3, sc4, sc5 = st.columns(5)
            with sc1:
                st.session_state["sec_summary"] = st.checkbox("Summary", True, key="cb_summary")
            with sc2:
                st.session_state["sec_skills"] = st.checkbox("Skills", True, key="cb_skills")
            with sc3:
                st.session_state["sec_experience"] = st.checkbox("Experience", True, key="cb_experience")
            with sc4:
                st.session_state["sec_education"] = st.checkbox("Education", True, key="cb_education")
            with sc5:
                st.session_state["sec_certs"] = st.checkbox("Certifications", False, key="cb_certs")

            c1, c2 = st.columns(2)
            with c1:
                if st.button("🔍 Analyze Fit (quick)", use_container_width=True):
                    with st.spinner("Analyzing..."):
                        if st.session_state.get("use_ai_tailor") and DEEPSEEK_KEY:
                            result = tailor_resume_ai(st.session_state.resume_text, job)
                            if result:
                                st.session_state.analysis = result
                            else:
                                rate, missing = tailor_resume_simple(st.session_state.resume_text, job)
                                st.session_state.analysis = f"**Keyword Match: {rate:.0f}%**\n\nMissing: {', '.join(missing[:10])}"
                                st.session_state.missing
                        else:
                            rate, missing = tailor_resume_simple(st.session_state.resume_text, job)
                            st.session_state.analysis = f"**Keyword Match: {rate:.0f}%**\n\nMissing: {', '.join(missing[:10])}"

            with c2:
                if st.button("✏️ Generate Tailored Resume (AI rewrite)", type="primary", use_container_width=True):
                    if not DEEPSEEK_KEY:
                        st.error("Need DeepSeek API key for this feature")
                    else:
                        with st.spinner("Rewriting resume for this job..."):
                            tailored = generate_tailored_resume(st.session_state.resume_text, job)
                            if tailored:
                                st.session_state.tailored_text = tailored
                                st.session_state.analysis = None
                            else:
                                st.error("AI call failed — try again")

            if "analysis" in st.session_state and st.session_state.analysis:
                st.divider()
                st.markdown(st.session_state.analysis)

                if "missing_skills" in st.session_state and st.session_state.missing_skills:
                    st.caption("💡 Missing keywords — add to resume if you have this experience:")
                    st.markdown("• " + "\n• ".join(st.session_state.missing_skills[:8]))

            if "tailored_text" in st.session_state and st.session_state.tailored_text:
                st.divider()
                # ── Resume Diff Preview ──
                if st.session_state.get("resume_text"):
                    with st.expander("🔍 See Changes (Original → Tailored)", expanded=False):
                        import difflib
                        orig_lines = st.session_state.resume_text.splitlines()
                        new_lines = st.session_state.tailored_text.splitlines()
                        
                        # Word-level inline diff
                        diff_html = []
                        sm = difflib.SequenceMatcher(None, orig_lines, new_lines)
                        for tag, i1, i2, j1, j2 in sm.get_opcodes():
                            if tag == 'equal':
                                for line in orig_lines[i1:i2]:
                                    diff_html.append(line)
                            elif tag == 'replace':
                                for old, new in zip(orig_lines[i1:i2], new_lines[j1:j2]):
                                    ws = difflib.SequenceMatcher(None, old.split(), new.split())
                                    parts = []
                                    for op, a1, a2, b1, b2 in ws.get_opcodes():
                                        if op == 'equal':
                                            parts.append(' '.join(old.split()[a1:a2]))
                                        elif op == 'delete':
                                            parts.append(f'<span style="background:#f8d7da;color:#721c24;text-decoration:line-through">{" ".join(old.split()[a1:a2])}</span>')
                                        elif op == 'insert':
                                            parts.append(f'<span style="background:#d4edda;color:#155724">{" ".join(new.split()[b1:b2])}</span>')
                                        elif op == 'replace':
                                            parts.append(f'<span style="background:#f8d7da;color:#721c24;text-decoration:line-through">{" ".join(old.split()[a1:a2])}</span>')
                                            parts.append(f'<span style="background:#d4edda;color:#155724">{" ".join(new.split()[b1:b2])}</span>')
                                    diff_html.append(' '.join(parts))
                            elif tag == 'delete':
                                for line in orig_lines[i1:i2]:
                                    diff_html.append(f'<span style="background:#f8d7da;color:#721c24;text-decoration:line-through">{line}</span>')
                            elif tag == 'insert':
                                for line in new_lines[j1:j2]:
                                    diff_html.append(f'<span style="background:#d4edda;color:#155724">{line}</span>')
                        
                        st.markdown('<br>'.join(diff_html[:150]), unsafe_allow_html=True)
                        st.caption("Red strikethrough = removed | Green = added | Inline = changed words")

                st.subheader("✏️ Tailored Resume — Edit Before Download")
                edited = st.text_area(
                    "Make any changes below. Add experience, tweak wording, or paste missing details.",
                    value=st.session_state.tailored_text,
                    height=400,
                    key="tailored_editor",
                )
                if edited != st.session_state.tailored_text:
                    st.session_state.tailored_text = edited
                    st.rerun()

                company_slug = re.sub(r'[^a-zA-Z0-9]', '_', job['company'])[:20]
                role_slug = job['title'][:30].replace(' ', '_')
                orig_name = st.session_state.get('resume_filename', 'resume')
                orig_base = re.sub(r'\.[^.]+$', '', orig_name)
                fname_base = f'{orig_base} - {company_slug} - {role_slug}'
                docx_buf = None
                if st.session_state.get("resume_bytes"):
                    docx_buf = build_tailored_docx(st.session_state.resume_bytes, st.session_state.tailored_text)
                if not docx_buf:
                    docx_buf = build_docx(st.session_state.tailored_text)
                st.download_button(
                    label="⬇️ Download .docx",
                    data=docx_buf,
                    file_name=f"{fname_base}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
    with tab3:
        st.subheader("Interview Questions")
        if st.button("🎲 Generate Questions", type="primary"):
            with st.spinner("Generating..."):
                if st.session_state.get("use_ai_questions") and DEEPSEEK_KEY:
                    qs = generate_questions_ai(job)
                else:
                    qs = None
                if not qs:
                    qs = generate_questions_simple(job)
                st.session_state.interview_qs = qs

        if "interview_qs" in st.session_state and st.session_state.interview_qs:
            for i, q in enumerate(st.session_state.interview_qs, 1):
                qtext = re.sub(r'^\d+[\.\)\-]+\s*', '', q)
                st.markdown(f"{i}. {qtext}")

    with tab4:
        st.subheader("Mock Interview")
        if "interview_qs" not in st.session_state or not st.session_state.interview_qs:
            st.warning("Generate questions first (❓ tab)")
        else:
            qs = st.session_state.interview_qs
            if "mock_idx" not in st.session_state:
                st.session_state.mock_idx = 0
                st.session_state.mock_feedback = []

            midx = st.session_state.mock_idx
            if midx >= len(qs):
                st.success("🎉 Interview complete!")
                for i, fb in enumerate(st.session_state.mock_feedback):
                    with st.expander(f"Q{i+1} feedback"):
                        st.markdown(fb or "*No feedback*")
                if st.button("🔄 Restart"):
                    st.session_state.mock_idx = 0
                    st.session_state.mock_feedback = []
                    st.rerun()
            else:
                st.progress(midx / len(qs), f"Question {midx+1}/{len(qs)}")
                # Strip any leading number prefix from question text (AI might still add one)
                qtext = re.sub(r'^\d+[\.\)\-]+\s*', '', qs[midx])
                st.markdown(f"##### Q{midx+1}: {qtext}")
                answer = st.text_area("Your answer:", key=f"ans_{midx}", height=120)

                c1, c2 = st.columns(2)
                with c1:
                    if st.button("✅ Submit", type="primary", use_container_width=True):
                        if answer.strip():
                            fb = evaluate_answer(qs[midx], answer) if DEEPSEEK_KEY else None
                            if not fb:
                                wc = len(answer.split())
                                fb = "Good detail and specificity." if wc > 30 else ("Decent, add examples." if wc > 10 else "Too short — elaborate.")
                            st.session_state.mock_feedback.append(fb)
                            st.session_state.mock_idx += 1
                            st.rerun()
                with c2:
                    if st.button("⏭ Skip", use_container_width=True):
                        st.session_state.mock_feedback.append("*Skipped*")
                        st.session_state.mock_idx += 1
                        st.rerun()

    with tab5:
        st.subheader("Questions to Ask the Interviewer")
        st.caption("Asking good questions shows preparation and interest. Pick 2-3 that matter most to you.")

        candidate_qs = [
            "What does a typical day or week look like in this role?",
            "How do you measure success for this position in the first 3-6 months?",
            "What are the biggest challenges the team is facing right now?",
            "How would you describe the team culture and working style?",
            "What opportunities for growth and learning does this role offer?",
            "How does the team handle technical decisions and disagreements?",
            "What's the onboarding process like for new team members?",
            "How does the company support work-life balance?",
            "Where do you see the company/team in the next 1-2 years?",
            "What do you enjoy most about working here?",
        ]
        # Add role-specific questions
        title = job.get('title', '').lower()
        if 'senior' in title or 'lead' in title:
            candidate_qs.append("What's your approach to technical leadership and mentoring on the team?")
        if 'manager' in title:
            candidate_qs.append("What's the team size and structure? How are reports distributed?")
        if 'engineer' in title or 'developer' in title:
            candidate_qs.append("What does the tech stack look like and how do you manage technical debt?")

        for i, q in enumerate(candidate_qs, 1):
            st.markdown(f"{i}. {q}")

        if DEEPSEEK_KEY:
            st.divider()
            if st.button("🤖 Generate Role-Specific Questions", key="gen_ask_them"):
                with st.spinner("Generating..."):
                    prompt = f"""Job: {job['title']} at {job['company']}
Description: {job.get('description','')[:1500]}

Generate 5 smart questions a candidate should ask the interviewer for this specific role and company. Make them thoughtful and specific. Return as numbered list."""
                    result = call_deepseek(prompt, "You are an experienced career coach helping candidates prepare for interviews.", max_tokens=300)
                    if result:
                        st.session_state.ask_them_qs = result
                        st.rerun()
            if "ask_them_qs" in st.session_state and st.session_state.ask_them_qs:
                st.markdown("##### AI-Generated Questions")
                st.markdown(st.session_state.ask_them_qs)


if __name__ == "__main__":
    main()
